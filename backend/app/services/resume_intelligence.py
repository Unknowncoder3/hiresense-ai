from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

SKILL_VOCABULARY = [
    "python", "sql", "excel", "power bi", "tableau", "pandas", "numpy",
    "scikit-learn", "tensorflow", "pytorch", "machine learning", "deep learning",
    "nlp", "aws", "azure", "gcp", "docker", "kubernetes", "fastapi", "flask",
    "django", "postgresql", "mysql", "mongodb", "redis", "java", "c++", "c",
    "javascript", "typescript", "react", "node.js", "git", "github", "mlflow",
    "statistics", "data analysis", "data visualization", "powerpoint", "jira",
    "a/b testing", "airflow", "spark", "hadoop", "communication"
]

ROLE_SKILLS = {
    "Data Analyst": ["SQL", "Python", "Power BI", "Excel", "Statistics", "Pandas"],
    "ML Engineer": ["Python", "Machine Learning", "Scikit-learn", "TensorFlow", "AWS"],
    "Product Analyst": ["SQL", "Excel", "A/B Testing", "Python", "Statistics"],
    "Backend Engineer": ["Python", "FastAPI", "PostgreSQL", "Docker", "Redis"],
    "BI Developer": ["SQL", "Power BI", "Excel", "DAX", "Data Visualization"],
}


def extract_text(filename: str, content: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        import io
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if ext == ".docx":
        import io
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()
    if ext == ".txt":
        return content.decode("utf-8", errors="ignore").strip()
    raise ValueError("Unsupported file type. Upload a PDF, DOCX, or TXT resume.")


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower()).strip()


def extract_skills(text: str) -> list[str]:
    normalized = normalize(text)
    found: list[str] = []
    for skill in sorted(SKILL_VOCABULARY, key=len, reverse=True):
        if re.search(rf"(?<![a-z0-9+#.]){re.escape(skill.lower())}(?![a-z0-9+#.])", normalized):
            pretty = skill.upper() if skill in {"sql", "aws", "gcp", "nlp", "c++", "c"} else skill.title()
            found.append(pretty)
    return list(dict.fromkeys(found))


def extract_years(text: str) -> float:
    patterns = [
        r"(\d+(?:\.\d+)?)\+?\s+years?\s+(?:of\s+)?experience",
        r"experience\s*[:\-]\s*(\d+(?:\.\d+)?)",
    ]
    values = []
    normalized = normalize(text)
    for pattern in patterns:
        values.extend(float(x) for x in re.findall(pattern, normalized))
    return max(values) if values else 0.0


def match_resume_to_job(resume_text: str, resume_skills: Iterable[str], job_title: str, candidate_experience: float = 0.0) -> dict:
    required = ROLE_SKILLS.get(job_title, ["Python", "SQL", "Communication"])
    skill_set = {s.lower() for s in resume_skills}
    matched = [skill for skill in required if skill.lower() in skill_set]
    missing = [skill for skill in required if skill.lower() not in skill_set]
    skill_score = len(matched) / len(required) * 70 if required else 0

    role_context = " ".join(required + [job_title, "candidate experience"])
    try:
        vectors = TfidfVectorizer(stop_words="english")
        matrix = vectors.fit_transform([resume_text, role_context])
        similarity = float(cosine_similarity(matrix[0:1], matrix[1:2])[0][0])
    except ValueError:
        similarity = 0.0

    experience_score = min(candidate_experience / 5, 1) * 20
    semantic_score = similarity * 10
    total = round(min(100, skill_score + experience_score + semantic_score), 1)

    strengths = matched[:5]
    gap = missing[:5]
    recommendation = (
        "Strong match. Prioritize this candidate for recruiter review."
        if total >= 85 else
        "Good potential match. Review the skill gaps before advancing."
        if total >= 70 else
        "Partial match. Consider the candidate only if the missing skills are trainable."
    )
    return {
        "score": total,
        "matched_skills": strengths,
        "skill_gaps": gap,
        "required_skills": required,
        "recommendation": recommendation,
        "breakdown": {
            "skill_match": round(skill_score, 1),
            "experience_fit": round(experience_score, 1),
            "semantic_similarity": round(semantic_score, 1),
        },
    }
